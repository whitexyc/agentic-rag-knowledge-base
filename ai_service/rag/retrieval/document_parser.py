"""统一文档解析层 — 多格式 → Markdown（module-064 / ADR-0014 WP1）

在整个 ingestion 链路中的位置：
  上传原始文件字节 → [DocumentParser 解析层] → Markdown → [清洗层] → 分块 → 嵌入

设计要点（ADR-0014 决策 1）：
  1. 格式识别读字节内容标记（anydoc.format_from_bytes 魔数探测），md/txt 无魔数
     用扩展名兜底——不靠扩展名一刀切（zip 类容器 docx/xlsx 靠 PK 头识别）。
  2. 主解析引擎 AnyDoc（firecrawl-anydoc，Firecrawl 开源 Rust 组件，14 格式统一
     GFM Markdown，零系统依赖无模型，Python 绑定释放 GIL）。
  3. PDF 回退：AnyDoc 失败 → PyMuPDF（存量 main.py:911 逻辑提取为可复用函数），
     保证存量 md/txt/pdf 行为兼容（存量零回归纪律）。
  4. docx/xlsx/csv 轻量回退（python-docx/openpyxl/csv 标准库）——AnyDoc 不可用
     时尽力解析；pptx/epub 无轻量回退则如实报"需 AnyDoc 解析引擎"。
  5. 错误变体映射用户中文提示（Unsupported/Malformed/Encrypted/ResourceLimit/
     MissingPart），上传端点直接返回提示。

诚实边界：
  - AnyDoc 安装失败 → _anydoc 为 None，PDF 走 PyMuPDF、docx/xlsx/csv 走轻量
    回退、pptx/epub 明确报错（如实降级，不假装支持）。
  - 扫描版 PDF（无文本层）不 OCR：AnyDoc/PyMuPDF 提取文本为空时由上层
    （image_pipeline WP4）如实返回"图片未解析"提示。
"""
import logging
import re
import unicodedata
from dataclasses import dataclass
from typing import Optional

logger = logging.getLogger(__name__)

# ── AnyDoc 延迟导入（可被测试 monkeypatch 替换） ─────────────────────────
try:
    import anydoc as anydoc  # noqa: F401  模块级全局，测试可 patch document_parser.anydoc
    from anydoc import (  # noqa: F401
        UnsupportedError,
        MalformedError,
        EncryptedError,
        MissingPartError,
        ResourceLimitError,
        ConvertError,
    )
    _ANYDOC_AVAILABLE = True
except Exception as _import_err:  # pragma: no cover - 环境探测
    anydoc = None  # type: ignore
    # AnyDoc 不可用时错误类全部归一为 Exception（调用方按类型捕获不命中即走通用路径）
    UnsupportedError = MalformedError = EncryptedError = Exception  # type: ignore
    MissingPartError = ResourceLimitError = ConvertError = Exception  # type: ignore
    _ANYDOC_AVAILABLE = False
    logger.warning("AnyDoc 解析引擎不可用（%s），PDF 走 PyMuPDF 回退、docx/xlsx/csv 走轻量回退", _import_err)

# 上传端允许的格式集合（前端 accept 同步，main.py upload 校验同源）
SUPPORTED_EXTENSIONS = (".md", ".txt", ".pdf", ".docx", ".xlsx", ".pptx", ".epub", ".csv")

# 扩展名 → 解析格式（md/txt 无魔数，AnyDoc format_from_bytes 返回 None，需扩展名兜底）
_EXTENSION_FORMATS = {
    ".md": "text",
    ".txt": "text",
    ".csv": "csv",
    ".json": "text",
    ".html": "html",
    ".htm": "html",
}

# AnyDoc 原生支持的格式（format_from_bytes / format_from_extension 返回值）
_ANYDOC_FORMATS = {"pdf", "docx", "xlsx", "pptx", "epub", "csv", "html"}


class DocumentParseError(Exception):
    """文档解析失败（携带面向用户的中文错误消息，上传端点直接透出）"""


@dataclass
class ParsedDocument:
    """解析结果：Markdown 文本 + 元信息

    Attributes:
        text: 解析出的 Markdown 文本
        format: 检测到的格式（pdf/docx/xlsx/pptx/epub/csv/text/...）
        engine: 实际解析引擎（anydoc/pymupdf/text/docx_fallback/xlsx_fallback/csv_fallback）
        page_count: PDF 页码（非 PDF 为 None；PDF 解析失败无法取得时 None）
    """
    text: str
    format: str
    engine: str
    page_count: Optional[int] = None


def _decode_text(data: bytes) -> str:
    """字节 → 文本：按 UTF-8 → GB18030 → Latin-1 依次尝试解码

    UTF-8 优先（现代文档默认），GB18030 兜底中文编码（Windows 生成的老 txt），
    Latin-1 最后兜底（任何字节都可解码，不抛错）。UTF-8 BOM 用 utf-8-sig 处理。
    """
    for enc in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def detect_format(data: bytes, filename: str = "") -> str:
    """格式识别：读字节内容标记，md/txt/csv 等无魔数格式用扩展名兜底

    优先 anydoc.format_from_bytes（魔数探测，docx/xlsx 靠 PK\x03\x04 头、
    pdf 靠 %PDF 头），返回 None 时按扩展名映射。最终未知抛 DocumentParseError。
    """
    if not data:
        raise DocumentParseError("上传文件为空")
    if anydoc is not None:
        try:
            fmt = anydoc.format_from_bytes(data)
            if fmt:
                return fmt
        except Exception as e:  # 探测失败不阻断，回退扩展名
            logger.warning("anydoc.format_from_bytes 探测失败: %s", e)
    ext = _extension_of(filename)
    if ext in _EXTENSION_FORMATS:
        return _EXTENSION_FORMATS[ext]
    if ext[1:] in _ANYDOC_FORMATS:  # 去点比对（".pdf" → "pdf"）
        return ext[1:]
    raise DocumentParseError(
        f"不支持的文件格式（{ext or '未知扩展名'}），请上传 "
        f"{'/'.join(SUPPORTED_EXTENSIONS)}"
    )


def _extension_of(filename: str) -> str:
    if not filename:
        return ""
    name = filename.lower().strip()
    # 取最后一个点后的后缀（含点），如 "a.PDF" → ".pdf"
    idx = name.rfind(".")
    return name[idx:] if idx >= 0 else ""


def _pdf_page_count(data: bytes) -> Optional[int]:
    """用 PyMuPDF 快速取 PDF 页数（不提取文本，仅 open 计数）"""
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        try:
            return doc.page_count
        finally:
            doc.close()
    except Exception:
        return None


def _parse_pdf_pymupdf(data: bytes) -> ParsedDocument:
    """PyMuPDF 回退解析（存量 main.py:911 逻辑提取为可复用函数）

    每页文本前加 `--- Page i/N ---` 分隔（与存量上传行为一致，零回归）。
    """
    try:
        import fitz
    except ImportError as e:
        raise DocumentParseError("PDF 解析库不可用，请安装 PyMuPDF") from e
    try:
        pdf_doc = fitz.open(stream=data, filetype="pdf")
        try:
            page_count = pdf_doc.page_count
            pages_text = []
            for i, page in enumerate(pdf_doc, start=1):
                text = page.get_text()
                pages_text.append(f"--- Page {i}/{page_count} ---\n{text}")
            full_text = "\n\n".join(pages_text)
            return ParsedDocument(text=full_text, format="pdf", engine="pymupdf",
                                  page_count=page_count)
        finally:
            pdf_doc.close()
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"PDF 解析失败: {e}") from e


def _parse_docx(data: bytes) -> ParsedDocument:
    """docx 轻量回退解析（python-docx，AnyDoc 不可用时尽力而为）

    段落 + 标题（Heading 样式 → 对应级 #）+ 表格 → Markdown。Best-effort，
    与 AnyDoc 的 GFM 输出可能存在排版差异（如实声明）。
    """
    try:
        from docx import Document as DocxDocument
        import io
        doc = DocxDocument(io.BytesIO(data))
        out: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name if para.style else "") or ""
            m = re.match(r"Heading\s*(\d+)", style_name)
            if m:
                level = min(int(m.group(1)), 6)
                out.append(f"{'#' * level} {text}")
            else:
                out.append(text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                n = rows[0].count("|") - 1
                rows.insert(1, "|" + "---|" * n)
                out.append("\n".join(rows))
        text = "\n\n".join(out)
        if not text.strip():
            raise DocumentParseError("Word 文档内容为空")
        return ParsedDocument(text=text, format="docx", engine="docx_fallback")
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"Word 文档解析失败: {e}") from e


def _rows_to_markdown(rows: list[list[str]]) -> str:
    """二维表 → GFM Markdown 表格（首行做表头 + 分隔行）"""
    md_rows = []
    for row in rows:
        cells = ["" if v is None else str(v).strip().replace("\n", " ") for v in row]
        if not any(cells):
            continue
        md_rows.append("| " + " | ".join(cells) + " |")
    if not md_rows:
        return ""
    n = md_rows[0].count("|") - 1
    md_rows.insert(1, "|" + "---|" * n)
    return "\n".join(md_rows)


def _parse_xlsx(data: bytes) -> ParsedDocument:
    """xlsx 轻量回退解析（openpyxl，AnyDoc 不可用时尽力而为）

    每 sheet → 一个 GFM 表格（首行表头）。Best-effort。
    """
    try:
        from openpyxl import load_workbook
        import io
        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        sheets_md = []
        for ws in wb.worksheets:
            rows = [list(row) for row in ws.iter_rows(values_only=True)]
            md = _rows_to_markdown(rows)
            if md:
                sheets_md.append(md)
        wb.close()
        text = "\n\n".join(sheets_md)
        if not text.strip():
            raise DocumentParseError("Excel 文档内容为空")
        return ParsedDocument(text=text, format="xlsx", engine="xlsx_fallback")
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"Excel 文档解析失败: {e}") from e


def _parse_csv(data: bytes) -> ParsedDocument:
    """csv 轻量回退解析（标准库 csv，AnyDoc 不可用时尽力而为）→ GFM 表格"""
    try:
        import csv
        import io
        text = _decode_text(data)
        reader = csv.reader(io.StringIO(text))
        rows = [r for r in reader if any(c.strip() for c in r)]
        md = _rows_to_markdown(rows)
        if not md.strip():
            raise DocumentParseError("CSV 文档内容为空")
        return ParsedDocument(text=md, format="csv", engine="csv_fallback")
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"CSV 文档解析失败: {e}") from e


def _map_error(e: Exception, fmt: str) -> str:
    """AnyDoc 错误变体 → 用户中文提示（WP1 错误映射）"""
    if isinstance(e, UnsupportedError):
        return f"不支持的文件格式（{fmt}），请上传 {'/'.join(SUPPORTED_EXTENSIONS)}"
    if isinstance(e, MalformedError):
        return f"文件已损坏或不是有效的 {fmt} 文件，无法解析"
    if isinstance(e, EncryptedError):
        return "文件已加密，无法解析（请先解密后再上传）"
    if isinstance(e, MissingPartError):
        return "文件不完整（缺少必要部分），无法解析"
    if isinstance(e, ResourceLimitError):
        return "文件过大或包含过多内容，无法解析"
    if isinstance(e, ConvertError):
        return f"文档解析失败，请检查 {fmt} 文件内容是否有效"
    return f"文档解析失败: {e}"


def parse_document(data: bytes, filename: str = "") -> ParsedDocument:
    """统一解析入口：字节 + 原始文件名 → Markdown 文本 + 元信息

    流程：
      1. 格式识别（读字节魔数，md/txt/csv 扩展名兜底）
      2. text（md/txt）→ 纯文本解码，不经 AnyDoc
      3. 其余格式 → AnyDoc 主解析；PDF/docx/xlsx/csv 失败走轻量回退；
         AnyDoc 整体不可用时 PDF 走 PyMuPDF、docx/xlsx/csv 走轻量回退、
         pptx/epub 明确报错（如实降级）

    Raises:
        DocumentParseError: 解析失败（message 为面向用户的中文提示）
    """
    if not data:
        raise DocumentParseError("上传文件为空")

    fmt = detect_format(data, filename)
    logger.info("parse_document: filename=%s, format=%s", filename or "(无)", fmt)

    # md/txt/html：纯文本解码，不经 AnyDoc（AnyDoc 不识别 text 格式）
    if fmt == "text":
        return ParsedDocument(text=_decode_text(data), format=fmt, engine="text")
    if fmt == "html":
        # html 交给 AnyDoc（转 Markdown）；AnyDoc 不可用时原样文本透出
        if anydoc is None:
            return ParsedDocument(text=_decode_text(data), format=fmt, engine="text")

    # AnyDoc 主解析路径
    if anydoc is not None:
        try:
            md = anydoc.to_markdown_bytes(data)
            if md is None:
                md = ""
            page_count = _pdf_page_count(data) if fmt == "pdf" else None
            return ParsedDocument(text=md, format=fmt, engine="anydoc",
                                  page_count=page_count)
        except Exception as e:
            # PDF 优先回退 PyMuPDF；docx/xlsx/csv 回退轻量解析器
            logger.warning("AnyDoc 解析失败（format=%s）: %s", fmt, e)
            if fmt == "pdf":
                return _parse_pdf_pymupdf(data)
            if fmt == "docx":
                return _parse_docx(data)
            if fmt == "xlsx":
                return _parse_xlsx(data)
            if fmt == "csv":
                return _parse_csv(data)
            raise DocumentParseError(_map_error(e, fmt)) from e

    # AnyDoc 不可用：分层回退（诚实降级）
    if fmt == "pdf":
        return _parse_pdf_pymupdf(data)
    if fmt == "docx":
        return _parse_docx(data)
    if fmt == "xlsx":
        return _parse_xlsx(data)
    if fmt == "csv":
        return _parse_csv(data)
    raise DocumentParseError(
        f"{fmt} 格式需要 AnyDoc 解析引擎（未安装），请上传 md/txt/pdf/docx/xlsx/csv"
    )
