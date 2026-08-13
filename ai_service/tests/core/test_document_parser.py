"""module-064 WP1 解析层测试（mock AnyDoc + 真实多格式，ADR-0014 决策 1）

覆盖：
- 格式识别读字节魔数（pdf/docx/xlsx），md/txt 无魔数用扩展名兜底
- md/txt 纯文本解码（UTF-8 / GBK 兜底）
- AnyDoc 主解析路径（mock）：pdf/page_count / docx
- 错误变体映射用户中文提示（Unsupported/Malformed/Encrypted/ResourceLimit）
- AnyDoc 不可用 → PDF 走 PyMuPDF、docx/xlsx/csv 走轻量回退、pptx/epub 明确报错
- 上传端点接线（扩展名校验 + 错误码）
"""
import asyncio
import io
import unittest.mock as mock

import pytest

from rag.retrieval import document_parser
from rag.retrieval.document_parser import (
    DocumentParseError,
    ParsedDocument,
    SUPPORTED_EXTENSIONS,
    detect_format,
    parse_document,
)


# ── 格式识别 ─────────────────────────────────────────────────────────────
def test_supported_extensions_cover_all():
    """上传端允许的 8 种格式齐全"""
    assert set(SUPPORTED_EXTENSIONS) == {
        ".md", ".txt", ".pdf", ".docx", ".xlsx", ".pptx", ".epub", ".csv",
    }


def test_detect_format_md_via_extension():
    """md/txt 无魔数：扩展名兜底 → text"""
    assert detect_format(b"hello", "a.md") == "text"
    assert detect_format(b"hello", "a.txt") == "text"


def test_detect_format_pdf_real_bytes():
    """PDF 魔数 %PDF 由 anydoc 探测"""
    import fitz
    buf = io.BytesIO()
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "probe")
    doc.save(buf)
    doc.close()
    assert detect_format(buf.getvalue(), "noext") == "pdf"


def test_detect_format_unknown_raises():
    """未知格式 → 明确 DocumentParseError"""
    with pytest.raises(DocumentParseError, match="不支持的文件格式"):
        detect_format(b"not a real doc at all", "a.xyz")


# ── 纯文本解析 ───────────────────────────────────────────────────────────
def test_parse_text_utf8():
    r = parse_document("你好 hello world".encode("utf-8"), "a.md")
    assert r.engine == "text"
    assert r.format == "text"
    assert "你好" in r.text


def test_parse_text_gbk_fallback():
    """GBK 编码老 txt：UTF-8 失败 → GB18030 兜底"""
    data = "中文内容".encode("gbk")
    r = parse_document(data, "a.txt")
    assert "中文内容" in r.text


def test_parse_empty_data():
    with pytest.raises(DocumentParseError, match="上传文件为空"):
        parse_document(b"", "a.md")


# ── AnyDoc 主解析（mock）────────────────────────────────────────────────
class FakeAnyDoc:
    """可编程假 anydoc：format_from_bytes / to_markdown_bytes"""

    def __init__(self, md="", raise_exc=None, detect=None):
        self._md = md
        self._raise = raise_exc
        self._detect = detect

    def format_from_bytes(self, data):
        return self._detect

    def to_markdown_bytes(self, data):
        if self._raise is not None:
            raise self._raise
        return self._md


def test_parse_pdf_anydoc(monkeypatch):
    """AnyDoc 主解析：pdf + page_count 透传"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(md="# PDF 标题\n\n正文", detect="pdf"))
    monkeypatch.setattr(document_parser, "_pdf_page_count", lambda data: 3)
    r = parse_document(b"%PDF-1.4 fake-bytes", "a.pdf")
    assert r.engine == "anydoc"
    assert r.format == "pdf"
    assert r.page_count == 3
    assert "# PDF 标题" in r.text


def test_parse_docx_anydoc(monkeypatch):
    """AnyDoc 主解析：docx → GFM Markdown"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(md="| A | B |\n| --- | --- |\n| 1 | 2 |", detect="docx"))
    r = parse_document(b"PK\x03\x04fake", "a.docx")
    assert r.engine == "anydoc"
    assert "| A | B |" in r.text


def test_parse_md_ignores_anydoc(monkeypatch):
    """md/txt 不走 AnyDoc（即使 anydoc 存在）——纯文本解码"""
    monkeypatch.setattr(document_parser, "anydoc", FakeAnyDoc(md="WRONG", detect=None))
    r = parse_document("真实文本".encode("utf-8"), "a.md")
    assert r.engine == "text"
    assert "真实文本" in r.text


# ── 错误变体映射 ────────────────────────────────────────────────────────
def test_error_unsupported_mapping(monkeypatch):
    """Unsupported → '不支持的文件格式'（pptx 无轻量回退 → 走映射）"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.UnsupportedError("unsupported"),
                                   detect="pptx"))
    with pytest.raises(DocumentParseError, match="不支持的文件格式"):
        parse_document(b"fake", "a.pptx")


def test_error_encrypted_mapping(monkeypatch):
    """Encrypted → '文件已加密'"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.EncryptedError("enc"),
                                   detect="pptx"))
    with pytest.raises(DocumentParseError, match="已加密"):
        parse_document(b"fake", "a.pptx")


def test_error_malformed_mapping(monkeypatch):
    """Malformed → '文件已损坏'"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.MalformedError("bad"),
                                   detect="pptx"))
    with pytest.raises(DocumentParseError, match="已损坏"):
        parse_document(b"fake", "a.pptx")


def test_pdf_anydoc_error_falls_back_to_pymupdf(monkeypatch):
    """PDF AnyDoc 失败 → PyMuPDF 回退（存量行为保留）"""
    monkeypatch.setattr(document_parser, "anydoc",
                        FakeAnyDoc(raise_exc=document_parser.ConvertError("conv"),
                                   detect="pdf"))
    fake = ParsedDocument(text="--- Page 1/1 ---\npyMuPDF 文本", format="pdf",
                          engine="pymupdf", page_count=1)
    monkeypatch.setattr(document_parser, "_parse_pdf_pymupdf", lambda data: fake)
    r = parse_document(b"%PDF-1.4", "a.pdf")
    assert r.engine == "pymupdf"
    assert "pyMuPDF 文本" in r.text


# ── AnyDoc 不可用 → 分层回退 ─────────────────────────────────────────────
def test_anydoc_unavailable_pdf_pymupdf(monkeypatch):
    monkeypatch.setattr(document_parser, "anydoc", None)
    fake = ParsedDocument(text="pymupdf text", format="pdf", engine="pymupdf", page_count=1)
    monkeypatch.setattr(document_parser, "_parse_pdf_pymupdf", lambda data: fake)
    r = parse_document(b"%PDF-1.4", "a.pdf")
    assert r.engine == "pymupdf"


def test_anydoc_unavailable_docx_fallback(monkeypatch):
    """AnyDoc 不可用 → python-docx 轻量回退（真实解析）"""
    monkeypatch.setattr(document_parser, "anydoc", None)
    from docx import Document as Docx
    doc = Docx()
    doc.add_heading("文档标题", 1)
    doc.add_paragraph("这是一段正文")
    buf = io.BytesIO()
    doc.save(buf)
    r = parse_document(buf.getvalue(), "a.docx")
    assert r.engine == "docx_fallback"
    assert "文档标题" in r.text
    assert "这是一段正文" in r.text


def test_anydoc_unavailable_xlsx_fallback(monkeypatch):
    monkeypatch.setattr(document_parser, "anydoc", None)
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws["A1"], ws["B1"] = "Name", "Value"
    ws["A2"], ws["B2"] = "Alice", 42
    buf = io.BytesIO()
    wb.save(buf)
    r = parse_document(buf.getvalue(), "a.xlsx")
    assert r.engine == "xlsx_fallback"
    assert "| Name | Value |" in r.text
    assert "| Alice | 42 |" in r.text


def test_anydoc_unavailable_pptx_clear_error(monkeypatch):
    """pptx 无轻量回退 → 明确报'需 AnyDoc 解析引擎'（诚实降级）"""
    monkeypatch.setattr(document_parser, "anydoc", None)
    with pytest.raises(DocumentParseError, match="需要 AnyDoc 解析引擎"):
        parse_document(b"PK\x03\x04", "a.pptx")


# ── 上传端点接线（AC 1.2/1.3）───────────────────────────────────────────
def test_upload_endpoint_unsupported_extension():
    """上传端点：非白名单扩展名 → code=1 明确提示"""
    import main as main_module

    async def run():
        with mock.patch("main.ingest_document") as mock_ingest:
            import httpx
            transport = httpx.ASGITransport(app=main_module.app, raise_app_exceptions=True)
            async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
                resp = await client.post(
                    "/ai/rag/documents/upload",
                    files={"file": ("a.xyz", b"whatever", "application/octet-stream")},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["code"] == 1
            assert "不支持的文件格式" in data["message"]
            mock_ingest.assert_not_called()

    asyncio.run(run())


def test_upload_endpoint_success_invokes_ingest():
    """上传端点：白名单格式 → 调用 ingest_document 并返回结果"""
    import main as main_module

    async def run():
        fake_result = {"id": 7, "title": "a", "chunks": 3, "duplicate": False,
                       "dup_kind": None, "page_count": None, "original_path": ""}
        with mock.patch("main.ingest_document",
                        new=mock.AsyncMock(return_value=fake_result)) as mock_ingest:
            import httpx
            transport = httpx.ASGITransport(app=main_module.app, raise_app_exceptions=True)
            async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
                resp = await client.post(
                    "/ai/rag/documents/upload",
                    files={"file": ("a.md", "# 标题\n\n正文".encode("utf-8"), "text/markdown")},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["code"] == 0
            assert data["data"]["id"] == 7
            mock_ingest.assert_awaited_once()

    asyncio.run(run())


def test_upload_endpoint_parse_error_mapped():
    """上传端点：解析失败 → code=3 + 中文提示"""
    import main as main_module

    async def run():
        with mock.patch("main.ingest_document",
                        new=mock.AsyncMock(
                            side_effect=DocumentParseError("文件已加密，无法解析"))):
            import httpx
            transport = httpx.ASGITransport(app=main_module.app, raise_app_exceptions=True)
            async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
                resp = await client.post(
                    "/ai/rag/documents/upload",
                    files={"file": ("a.pdf", b"%PDF-1.4 fake", "application/pdf")},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["code"] == 3
            assert "已加密" in data["message"]

    asyncio.run(run())
